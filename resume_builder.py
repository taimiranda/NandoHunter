"""Utilities for persisting generated resume artifacts."""

import re
from pathlib import Path


def _append_markdown_runs(paragraph, text: str) -> None:
	"""Append paragraph runs honoring simple **bold** markdown segments."""
	parts = re.split(r"(\*\*.*?\*\*)", text)
	for part in parts:
		if not part:
			continue
		if part.startswith("**") and part.endswith("**") and len(part) >= 4:
			run = paragraph.add_run(part[2:-2])
			run.bold = True
		else:
			paragraph.add_run(part)


def markdown_to_docx(markdown_text: str, output_path: str) -> str:
	"""Convert markdown into a clean .docx file."""
	from docx import Document

	doc = Document()

	lines = (markdown_text or "").splitlines()
	for raw_line in lines:
		line = raw_line.strip()
		if not line:
			doc.add_paragraph("")
			continue

		if line.startswith("# "):
			paragraph = doc.add_paragraph(style="Title")
			_append_markdown_runs(paragraph, line[2:].strip())
			continue

		if line.startswith("## "):
			paragraph = doc.add_paragraph(style="Heading 1")
			_append_markdown_runs(paragraph, line[3:].strip())
			continue

		if line.startswith("### "):
			paragraph = doc.add_paragraph(style="Heading 2")
			_append_markdown_runs(paragraph, line[4:].strip())
			continue

		if line.startswith("- "):
			paragraph = doc.add_paragraph(style="List Bullet")
			_append_markdown_runs(paragraph, line[2:].strip())
			continue

		paragraph = doc.add_paragraph(style="Normal")
		_append_markdown_runs(paragraph, line)

	out_path = Path(output_path)
	out_path.parent.mkdir(parents=True, exist_ok=True)
	doc.save(str(out_path))
	return str(out_path)


def save_markdown_files(
	job_id: int,
	bullets: str,
	cover_letter: str,
	job_title: str,
	company: str,
) -> tuple[str, str]:
	"""Save generated resume and cover letter to markdown files."""
	output_dir = Path("./output")
	output_dir.mkdir(parents=True, exist_ok=True)

	filename = re.sub(r"[^\w\s-]", "", f"{company}_{job_title}").strip()
	if not filename:
		filename = f"job_{job_id}"

	resume_path = output_dir / f"resume_{filename}.md"
	cover_letter_path = output_dir / f"cover_{filename}.md"

	resume_path.write_text(bullets or "", encoding="utf-8")
	cover_letter_path.write_text(cover_letter or "", encoding="utf-8")

	return str(resume_path), str(cover_letter_path)


def save_all_formats(
	job_id: int,
	resume_md: str,
	cover_letter_md: str,
	job_title: str,
	company: str,
) -> dict:
	"""Save markdown and docx versions for both resume and cover letter."""
	output_dir = Path("./output")
	output_dir.mkdir(parents=True, exist_ok=True)

	safe_name = re.sub(r"[^\w-]", "_", f"{company}_{job_title}")
	if not safe_name:
		safe_name = f"job_{job_id}"

	resume_md_path = output_dir / f"resume_{safe_name}.md"
	resume_docx_path = output_dir / f"resume_{safe_name}.docx"
	cover_md_path = output_dir / f"cover_{safe_name}.md"
	cover_docx_path = output_dir / f"cover_{safe_name}.docx"

	resume_md_path.write_text(resume_md or "", encoding="utf-8")
	cover_md_path.write_text(cover_letter_md or "", encoding="utf-8")

	markdown_to_docx(resume_md or "", str(resume_docx_path))
	markdown_to_docx(cover_letter_md or "", str(cover_docx_path))

	return {
		"resume_md_path": str(resume_md_path),
		"resume_docx_path": str(resume_docx_path),
		"cover_md_path": str(cover_md_path),
		"cover_docx_path": str(cover_docx_path),
	}
